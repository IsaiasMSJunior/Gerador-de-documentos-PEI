import streamlit as st
import json
import os
import pandas as pd
import calendar
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from datetime import datetime
from io import BytesIO
from openpyxl import load_workbook
from openpyxl.styles import PatternFill
import mysql.connector

conn = None

# Configurações de conexão
host = "localhost"  # Endereço do servidor MySQL
user = "gerador1_ARQUIVOS"  # Nome de usuário
password = "mudar123"  # Senha
database = "teste"  # Nome do banco de dados

# Cria a conexão
try:
    conn = mysql.connector.connect(host=host, user=user, password=password, database=database)

    if conn.is_connected():
        print("Conexão bem-sucedida!")
        # Código para interagir com o banco de dados
    else:
        print("Erro ao conectar ao banco de dados.")

except mysql.connector.Error as err:
    print(f"Erro: {err}")

finally:
    if conn.is_connected():
        conn.close()
        print("Conexão fechada.")

st.set_page_config(page_title="Agenda Escolar", layout="wide")

# --- Helpers e Configurações Gerais ---

map_hor = {
    "1ª": "7:00–7:50", "2ª": "7:50–8:40", "3ª": "8:40–9:30",
    "4ª": "9:50–10:40", "5ª": "10:40–11:30", "6ª": "12:20–13:10", "7ª": "13:10–14:00"
}
meses = ["Janeiro","Fevereiro","Março","Abril","Maio","Junho",
         "Julho","Agosto","Setembro","Outubro","Novembro","Dezembro"]
ano_planej = 2025

def carregar_json(nome):
    if os.path.exists(nome):
        with open(nome, 'r', encoding='utf-8') as f:
            return json.load(f)
    return {}

def salvar_json(nome, conteudo):
    with open(nome, 'w', encoding='utf-8') as f:
        json.dump(conteudo, f, ensure_ascii=False, indent=2)

def extrai_serie(turma: str) -> str:
    return turma[:-1]

def set_border(par: Paragraph):
    p = par._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bd = OxmlElement('w:bottom')
    bd.set(qn('w:val'),'single'); bd.set(qn('w:sz'),'4')
    bd.set(qn('w:space'),'1'); bd.set(qn('w:color'),'auto')
    pBdr.append(bd); pPr.append(pBdr)

def insert_after(par: Paragraph, text='') -> Paragraph:
    new_p = OxmlElement('w:p'); par._p.addnext(new_p)
    para = Paragraph(new_p, par._parent)
    if text: para.add_run(text)
    return para

# --- Funções de geração de documentos ---

def gerar_agenda_template(entries, df_bank, professor, semana, bimestre, cores_turmas):
    wb = load_workbook("agenda_modelo.xlsx")
    ws = wb.active
    ws["B1"] = professor
    ws["E1"] = semana
    row_map = {"1ª":4, "2ª":6, "3ª":8, "4ª":12, "5ª":14, "6ª":18, "7ª":20}
    col_map = {"Segunda":"C", "Terça":"D", "Quarta":"E", "Quinta":"F", "Sexta":"G"}
    for e in entries:
        col, row = col_map[e["dia"]], row_map[e["aula"]]
        ws[f"{col}{row}"] = f"{e['turma']} – {e['disciplina']}"
        fill = PatternFill(
            start_color=cores_turmas.get(e["turma"], "#FFFFFF").lstrip("#"),
            end_color=cores_turmas.get(e["turma"], "#FFFFFF").lstrip("#"),
            fill_type="solid"
        )
        ws[f"{col}{row}"].fill = fill
        ws[f"{col}{row+1}"] = (
            f"Aula {e['num']} – " +
            df_bank.loc[
                (df_bank["DISCIPLINA"]==e["disciplina"]) &
                (df_bank["ANO/SÉRIE"]==extrai_serie(e["turma"])) &
                (df_bank["BIMESTRE"]==bimestre) &
                (df_bank["Nº da aula"]==e["num"])
            ]["TÍTULO DA AULA"].iloc[0]
            if not df_bank.empty else ""
        )
        ws[f"{col}{row+1}"].fill = fill
    out = BytesIO(); wb.save(out); out.seek(0)
    return out

def gerar_plano_template(entries, df_bank, professor, semana, bimestre, turma,
                         metodologias, recursos, criterios, modelo="modelo_plano.docx"):
    doc = Document(modelo)
    header_disciplinas = ", ".join(sorted({e['disciplina'] for e in entries}))
    total_aulas = str(len(entries))
    # Cabeçalho
    for p in doc.paragraphs:
        p.text = (p.text
                  .replace("ppp", professor)
                  .replace("ttt", turma)
                  .replace("sss", semana)
                  .replace("ddd", header_disciplinas)
                  .replace("nnn", total_aulas))
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.text = (p.text
                              .replace("ppp", professor)
                              .replace("ttt", turma)
                              .replace("sss", semana)
                              .replace("ddd", header_disciplinas)
                              .replace("nnn", total_aulas))
    # Blocos de aula + extras
    for p in doc.paragraphs:
        if p.text.strip() == "ccc":
            p.text = ""; last = p
            b0 = insert_after(last); set_border(b0); last = b0
            last = insert_after(last, "")
            for e in entries:
                sub = df_bank.loc[
                    (df_bank["DISCIPLINA"]==e["disciplina"]) &
                    (df_bank["ANO/SÉRIE"]==extrai_serie(turma)) &
                    (df_bank["BIMESTRE"]==bimestre) &
                    (df_bank["Nº da aula"]==e["num"])
                ]
                titulo = sub["TÍTULO DA AULA"].iloc[0] if not sub.empty else ""
                hab    = sub["HABILIDADE"].iloc[0]        if not sub.empty else ""
                cnt    = sub["CONTEÚDO"].iloc[0]         if not sub.empty else ""
                pa = insert_after(last, f"Aula {e['num']} – {titulo}"); pa.runs[0].bold=True; last=pa
                last = insert_after(last, "")
                ph = insert_after(last); rh=ph.add_run("Habilidade: "); rh.underline=True; ph.add_run(hab); last=ph
                last = insert_after(last, "")
                pc = insert_after(last); rc=pc.add_run("Conteúdo: "); rc.underline=True; pc.add_run(cnt); last=pc
                last = insert_after(last, "")
                b1 = insert_after(last); set_border(b1); last=b1
                last = insert_after(last, "")
            if metodologias:
                pm = insert_after(last); pm.add_run("Metodologia:").bold=True; last=pm
                for m in metodologias: last=insert_after(last, f"• {m}")
                last=insert_after(last, "")
            if recursos:
                pr = insert_after(last); pr.add_run("Recursos:").bold=True; last=pr
                for r in recursos: last=insert_after(last, f"• {r}")
                last=insert_after(last, "")
            if criterios:
                pc2 = insert_after(last); pc2.add_run("Critérios de Avaliação:").bold=True; last=pc2
                for c in criterios: last=insert_after(last, f"• {c}")
            break
    out = BytesIO(); doc.save(out); out.seek(0)
    return out

def gerar_guia_template(professor, turma, disciplina, bimestre, inicio, fim,
                        qtd_bimestre, qtd_semanal, metodologias, criterios,
                        df_bank, modelo="modelo_guia.docx"):
    doc = Document(modelo)
    reps = {
        'ppp': professor,
        'ttt': turma,
        'bbb': bimestre,
        'iii': inicio.strftime('%d/%m/%Y'),
        'fff': fim.strftime('%d/%m/%Y'),
        'qqq': str(qtd_bimestre),
        'sss': str(qtd_semanal),
        'mmm': ", ".join(metodologias),
        'ccc': ", ".join(criterios),
        'ddd': disciplina
    }
    for p in doc.paragraphs:
        for k,v in reps.items():
            if k in p.text: p.text = p.text.replace(k,v)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for k,v in reps.items():
                        if k in p.text: p.text = p.text.replace(k,v)
    # habilidades e objetos únicos
    mask = (
        (df_bank["DISCIPLINA"]==disciplina)&
        (df_bank["ANO/SÉRIE"]==extrai_serie(turma))&
        (df_bank["BIMESTRE"]==bimestre)
    )
    habs = df_bank.loc[mask, "HABILIDADE"].dropna().astype(str).tolist()
    objs = df_bank.loc[mask, "OBJETO DE CONHECIMENTO"].dropna().astype(str).tolist()
    unique_habs, unique_objs = [], []
    for h in habs:
        if h not in unique_habs: unique_habs.append(h)
    for o in objs:
        if o not in unique_objs: unique_objs.append(o)
    # substitui hhh e ooo
    for p in doc.paragraphs:
        if 'hhh' in p.text: p.text = "\n".join(unique_habs)
        if 'ooo' in p.text: p.text = "\n".join(unique_objs)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if 'hhh' in p.text: p.text = "\n".join(unique_habs)
                    if 'ooo' in p.text: p.text = "\n".join(unique_objs)
    out = BytesIO(); doc.save(out); out.seek(0)
    return out

def gerar_planejamento_template(professor, disciplina, turma, bimestre,
                                grupos, df_bank, modelo="modelo_planejamento.docx"):
    doc = Document(modelo)
    # cabeçalho
    hdr = {'ppp':professor,'ddd':disciplina,'ttt':turma,'bbb':bimestre}
    for p in doc.paragraphs:
        for k,v in hdr.items():
            if k in p.text: p.text = p.text.replace(k,v)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for k,v in hdr.items():
                        if k in p.text: p.text = p.text.replace(k,v)
    # corpo por grupo
    for grp in grupos:
        p0 = doc.add_paragraph(); set_border(p0)
        doc.add_paragraph()
        doc.add_paragraph(f"Semana: {grp['semana']}")
        doc.add_paragraph(f"Aulas previstas: {grp['prev']}")
        doc.add_paragraph("Aulas dadas:")
        doc.add_paragraph("Aulas do material digital:")
        for n in grp['nums']:
            mask = (
                (df_bank["DISCIPLINA"]==disciplina)&
                (df_bank["ANO/SÉRIE"]==extrai_serie(turma))&
                (df_bank["BIMESTRE"]==bimestre)&
                (df_bank["Nº da aula"]==n)
            )
            titles = df_bank.loc[mask,"TÍTULO DA AULA"].dropna().astype(str).tolist()
            title = titles[0] if titles else ""
            doc.add_paragraph(f"Aula {n} – {title}")
        doc.add_paragraph("Objetivos:")
        objs = df_bank.loc[mask,"OBJETIVOS"].dropna().astype(str).tolist()
        seen = []
        for o in objs:
            if o not in seen:
                seen.append(o); doc.add_paragraph(o)
        doc.add_paragraph("Habilidade:")
        habs = df_bank.loc[mask,"HABILIDADE"].dropna().astype(str).tolist()
        seen = []
        for h in habs:
            if h not in seen:
                seen.append(h); doc.add_paragraph(h)
        doc.add_paragraph(f"Metodologia: {', '.join(grp['met'])}")
        doc.add_paragraph(f"Critérios de avaliação: {', '.join(grp['crit'])}")
        doc.add_paragraph()
    out = BytesIO(); doc.save(out); out.seek(0)
    return out

# --- Inicialização de estados ---

if "extras" not in st.session_state:
    extras = carregar_json("extras.json") or {}
    extras.setdefault("metodologia",[])
    extras.setdefault("recursos",[])
    extras.setdefault("criterios",[])
    st.session_state.extras = extras

# Sessão de dados
if "pagina" not in st.session_state:
    st.session_state.pagina = "Cadastro de Professor"
if "professores" not in st.session_state:
    st.session_state.professores = carregar_json("professores.json") or []
if "turmas" not in st.session_state:
    st.session_state.turmas = carregar_json("turmas.json") or {}
if "horarios" not in st.session_state:
    st.session_state.horarios = carregar_json("horarios.json") or []

# --- Sidebar ---

pages = [
    "Cadastro de Professor","Cadastro de Turmas","Cadastro de Horário",
    "Gerar Agenda e Plano","Cadastro Extras","Gerar Guia",
    "Gerar Planejamento Bimestral"
]
for p in pages:
    if st.sidebar.button(p, use_container_width=True):
        st.session_state.pagina = p
    st.sidebar.markdown("\n")

# --- Páginas ---

# 1. Cadastro de Professor
if st.session_state.pagina == "Cadastro de Professor":
    st.header("Cadastro de Professor")
    nome = st.text_input("Nome")
    disciplinas = st.multiselect(
        "Disciplina(s)",
        ["Arte","Ciências","Ed. Física","Ed. Financeira","Geografia","História",
         "Português","Inglês","Matemática","PV","Redação","Tecnologia","OE Port","OE Mat"]
    )
    if st.button("Salvar Professor"):
        st.session_state.professores.append({"nome":nome,"disciplinas":disciplinas})
        salvar_json("professores.json", st.session_state.professores)
        st.success("Professor salvo!")
    for p in st.session_state.professores:
        st.write(f"{p['nome']} — {', '.join(p['disciplinas'])}")

# 2. Cadastro de Turmas
elif st.session_state.pagina == "Cadastro de Turmas":
    st.header("Cadastro de Turmas")
    saved = st.session_state.turmas
    default_s = sorted({t[:-1] for t in saved.keys()})
    default_seg = []
    if any(s in ["6º","7º","8º","9º"] for s in default_s): default_seg.append("Ensino Fundamental")
    if any(s in ["1º","2º","3º"] for s in default_s): default_seg.append("Ensino Médio")
    segmento = st.multiselect("Segmento(s)", ["Ensino Fundamental","Ensino Médio"], default=default_seg)
    anos = []
    if "Ensino Fundamental" in segmento: anos+=["6º","7º","8º","9º"]
    if "Ensino Médio" in segmento: anos+=["1º","2º","3º"]
    series = st.multiselect("Ano/Série", anos, default=default_s)
    turma_map = {
        "6º":["6ºA","6ºB","6ºC","6ºD"],"7º":["7ºA","7ºB","7ºC"],
        "8º":["8ºA","8ºB","8ºC","8ºD"],"9º":["9ºA","9ºB","9ºC","9ºD"],
        "1º":["1ºA","1ºB","1ºC","1ºD","1ºE"],
        "2º":["2ºA ADM","2ºB ADM","2ºC"],"3º":["3ºA","3ºA ADM","3ºB ADM","3ºB LOG"]
    }
    op = sum((turma_map.get(s,[]) for s in series), [])
    sel = st.multiselect("Turma(s)", op, default=list(saved.keys()), key="sel_turmas")
    cores = {
        t: st.color_picker(f"Cor {t}", value=saved.get(t,"#FFFFFF"), key=f"cor_{t}")
        for t in sel
    }
    if st.button("Salvar Turmas"):
        st.session_state.turmas = cores
        salvar_json("turmas.json", st.session_state.turmas)
        st.success("Turmas salvas!")

# 3. Cadastro de Horário
elif st.session_state.pagina == "Cadastro de Horário":
    st.header("Cadastro de Horário")
    if st.button("Adicionar Linha"):
        st.session_state.horarios.append({'turma':None,'disciplina':None,'dia':None,'aula':None})
    for i, itm in enumerate(st.session_state.horarios):
        cols = st.columns(6)
        turmas = list(st.session_state.turmas.keys())
        discs  = sorted({d for p in st.session_state.professores for d in p["disciplinas"]})
        dias   = ["Segunda","Terça","Quarta","Quinta","Sexta"]
        aulas  = list(map_hor.keys())
        itm['turma']      = cols[0].selectbox("Turma", turmas,
                                index=turmas.index(itm.get('turma')) if itm.get('turma') in turmas else 0,
                                key=f"turma_{i}")
        itm['disciplina'] = cols[1].selectbox("Disciplina", discs,
                                index=discs.index(itm.get('disciplina')) if itm.get('disciplina') in discs else 0,
                                key=f"disc_{i}")
        itm['dia']        = cols[2].selectbox("Dia", dias,
                                index=dias.index(itm.get('dia')) if itm.get('dia') in dias else 0,
                                key=f"dia_{i}")
        itm['aula']       = cols[3].selectbox("Aula", aulas,
                                index=aulas.index(itm.get('aula')) if itm.get('aula') in aulas else 0,
                                key=f"aula_{i}")
        cols[4].text_input("Horário", map_hor.get(itm['aula'],""), disabled=True, key=f"hor_{i}")
        if cols[5].button("X", key=f"rm_{i}"):
            st.session_state.horarios.pop(i)
            break
    if st.button("Salvar Horários"):
        salvar_json("horarios.json", st.session_state.horarios)
        st.success("Horários salvos!")
    if st.session_state.horarios:
        st.dataframe(pd.DataFrame(st.session_state.horarios).sort_values("dia"))

# 4. Gerar Agenda e Plano
elif st.session_state.pagina == "Gerar Agenda e Plano":
    st.header("Gerar Agenda e Plano")
    st.markdown("<style>[role='tab']{font-size:18px!important;padding:0.75rem 1.5rem!important;}</style>",
                unsafe_allow_html=True)
    if not st.session_state.horarios:
        st.warning("Cadastre horários primeiro.")
    else:
        df_bank = pd.read_excel("ES_banco.xlsx", header=0)
        prof = st.selectbox("Professor(a)", [p["nome"] for p in st.session_state.professores])
        bim = st.selectbox("Bimestre", ["1º","2º","3º","4º"])
        meses_lista = meses
        mes_nome  = st.selectbox("Mês", meses_lista)
        semanas = [
            f"{w[0].strftime('%d/%m')} – {w[-1].strftime('%d/%m')}"
            for w in calendar.Calendar().monthdatescalendar(datetime.now().year, meses_lista.index(mes_nome)+1)
            if w[0].month == meses_lista.index(mes_nome)+1
        ]
        sem_sel = st.selectbox("Semana", semanas)
        turma_idx = {}
        for idx, itm in enumerate(st.session_state.horarios):
            turma_idx.setdefault(itm['turma'], []).append(idx)
        entries = []
        tabs = st.tabs(list(turma_idx.keys()))
        for tab, turma in zip(tabs, turma_idx.keys()):
            with tab:
                st.subheader(f"Turma {turma}")
                met_sel = st.multiselect("Metodologia", st.session_state.extras["metodologia"], key=f"met_sel_{turma}")
                rec_sel = st.multiselect("Recursos", st.session_state.extras["recursos"], key=f"rec_sel_{turma}")
                crit_sel = st.multiselect("Critérios de Avaliação", st.session_state.extras["criterios"], key=f"crit_sel_{turma}")
                for idx in turma_idx[turma]:
                    h = st.session_state.horarios[idx]
                    st.markdown(f"**{turma} | {h['disciplina']} | {h['dia']} | {h['aula']}**")
                    mask = (
                        (df_bank["DISCIPLINA"]==h['disciplina']) &
                        (df_bank["ANO/SÉRIE"]==extrai_serie(turma)) &
                        (df_bank["BIMESTRE"]==bim)
                    )
                    opts = df_bank.loc[mask, "Nº da aula"].dropna().unique()
                    nums = sorted({int(x) for x in opts})
                    num = st.selectbox("Nº da aula", nums, key=f"num_{turma}_{idx}")
                    entries.append({**h, "num": num})
                if st.button("Gerar Plano", key=f"gera_plano_{turma}"):
                    arq = gerar_plano_template(
                        [e for e in entries if e["turma"]==turma],
                        df_bank, prof, sem_sel, bim, turma,
                        metodologias=met_sel, recursos=rec_sel, criterios=crit_sel
                    )
                    st.download_button(f"Download Plano {turma}", data=arq,
                                       file_name=f"plano_{turma}.docx",
                                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        if st.button("Gerar Agenda"):
            ag = gerar_agenda_template(entries, df_bank, prof, sem_sel, bim, st.session_state.turmas)
            st.download_button("Download Agenda", data=ag,
                               file_name="agenda_preenchida.xlsx",
                               mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

# 5. Cadastro Extras
elif st.session_state.pagina == "Cadastro Extras":
    st.header("Cadastro Extras")
    tab1, tab2, tab3 = st.tabs(["Metodologia","Recursos","Critérios de Avaliação"])
    with tab1:
        st.text_input("Metodologia", key="input_met")
        st.button("Inserir Metodologia", on_click=lambda: st.session_state.extras["metodologia"].append(st.session_state.input_met) or salvar_json("extras.json", st.session_state.extras) or st.session_state.update(input_met=""))
        for i, item in enumerate(st.session_state.extras["metodologia"]):
            c1,c2 = st.columns([0.9,0.1])
            c1.write(f"- {item}"); c2.button("X", key=f"del_met_{i}", on_click=lambda i=i: st.session_state.extras["metodologia"].pop(i) or salvar_json("extras.json", st.session_state.extras))
    with tab2:
        st.text_input("Recursos", key="input_rec")
        st.button("Inserir Recursos", on_click=lambda: st.session_state.extras["recursos"].append(st.session_state.input_rec) or salvar_json("extras.json", st.session_state.extras) or st.session_state.update(input_rec=""))
        for i, item in enumerate(st.session_state.extras["recursos"]):
            c1,c2 = st.columns([0.9,0.1])
            c1.write(f"- {item}"); c2.button("X", key=f"del_rec_{i}", on_click=lambda i=i: st.session_state.extras["recursos"].pop(i) or salvar_json("extras.json", st.session_state.extras))
    with tab3:
        st.text_input("Critério de Avaliação", key="input_crit")
        st.button("Inserir Critério", on_click=lambda: st.session_state.extras["criterios"].append(st.session_state.input_crit) or salvar_json("extras.json", st.session_state.extras) or st.session_state.update(input_crit=""))
        for i, item in enumerate(st.session_state.extras["criterios"]):
            c1,c2 = st.columns([0.9,0.1])
            c1.write(f"- {item}"); c2.button("X", key=f"del_crit_{i}", on_click=lambda i=i: st.session_state.extras["criterios"].pop(i) or salvar_json("extras.json", st.session_state.extras))

# 6. Gerar Guia
elif st.session_state.pagina == "Gerar Guia":
    st.header("Gerar Guia")
    if not st.session_state.horarios:
        st.warning("Cadastre horários primeiro.")
    else:
        df_bank = pd.read_excel("ES_banco.xlsx", header=0)
        prof = st.selectbox("Professor(a)", [p["nome"] for p in st.session_state.professores])
        bim = st.selectbox("Bimestre", ["1º","2º","3º","4º"])
        inicio = st.date_input("Início"); fim = st.date_input("Fim")
        turmas = sorted({h["turma"] for h in st.session_state.horarios})
        tabs = st.tabs(turmas)
        for tab, turma in zip(tabs, turmas):
            with tab:
                st.subheader(f"Turma {turma}")
                disc_opts = sorted({h["disciplina"] for h in st.session_state.horarios if h["turma"]==turma})
                disciplina = st.selectbox("Disciplina", disc_opts, key=f"disc_g_{turma}")
                qtd_bim = st.number_input("Qtd. de aulas no Bimestre", min_value=1, key=f"q_bim_{turma}")
                qtd_sem = st.number_input("Qtd. de aulas semanais", min_value=1, key=f"q_sem_{turma}")
                met_sel = st.multiselect("Metodologia de Ensino", st.session_state.extras["metodologia"], key=f"met_g_{turma}")
                crit_sel = st.multiselect("Como serei Avaliado", st.session_state.extras["criterios"], key=f"crit_g_{turma}")
                if st.button("Gerar Guia", key=f"gera_guia_{turma}"):
                    arq = gerar_guia_template(prof, turma, disciplina, bim, inicio, fim, qtd_bim, qtd_sem, met_sel, crit_sel, df_bank)
                    st.download_button(f"Download Guia {turma}", data=arq, file_name=f"guia_{turma}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# 7. Gerar Planejamento Bimestral
elif st.session_state.pagina == "Gerar Planejamento Bimestral":
    st.header("Gerar Planejamento Bimestral")
    if not st.session_state.horarios:
        st.warning("Cadastre horários primeiro.")
    else:
        df_bank = pd.read_excel("ES_banco.xlsx", header=0)
        prof = st.selectbox("Professor(a)", [p["nome"] for p in st.session_state.professores])
        bim = st.selectbox("Bimestre", ["1º","2º","3º","4º"])
        turmas = sorted({h["turma"] for h in st.session_state.horarios})
        tabs = st.tabs(turmas)
        for tab, turma in zip(tabs, turmas):
            with tab:
                st.subheader(f"Turma {turma}")
                cnt_key = f"count_{turma}"
                if cnt_key not in st.session_state: st.session_state[cnt_key] = 1
                grupos = []
                for i in range(st.session_state[cnt_key]):
                    with st.expander(f"Planejamento {i+1}", expanded=True):
                        mes = st.selectbox("Mês", meses, key=f"plan_mes_{turma}_{i}")
                        mi = meses.index(mes)+1
                        semanas = [
                            f"{w[0].strftime('%d/%m')} – {w[4].strftime('%d/%m')}"
                            for w in calendar.Calendar().monthdatescalendar(ano_planej, mi)
                            if all(d.month==mi for d in w[:5])
                        ]
                        semana = st.selectbox("Semana", semanas, key=f"plan_sem_{turma}_{i}")
                        prev = st.number_input("Aulas previstas", min_value=1, key=f"plan_prev_{turma}_{i}")
                        nums_opts = sorted(int(k.replace("ª","")) for k in map_hor.keys())
                        nums = st.multiselect("Nº das aulas", nums_opts, key=f"plan_nums_{turma}_{i}")
                        met = st.multiselect("Metodologias", st.session_state.extras["metodologia"], key=f"plan_met_{turma}_{i}")
                        crit = st.multiselect("Critérios de avaliação", st.session_state.extras["criterios"], key=f"plan_crit_{turma}_{i}")
                        grupos.append({"semana":semana,"prev":prev,"nums":nums,"met":met,"crit":crit})
                if st.button("Adicionar", key=f"add_plan_{turma}"):
                    st.session_state[cnt_key] += 1
                if st.button("Gerar Planejamento", key=f"gera_plan_{turma}"):
                    disc_set = sorted({h["disciplina"] for h in st.session_state.horarios if h["turma"]==turma})
                    disciplina = ", ".join(disc_set)
                    arq = gerar_planejamento_template(prof, disciplina, turma, bim, grupos, df_bank)
                    st.download_button(f"Download Planejamento {turma}", data=arq,
                                       file_name=f"planejamento_{turma}.docx",
                                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
