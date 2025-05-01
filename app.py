import streamlit as st
st.set_page_config(page_title="Agenda Escolar", layout="wide")
st.write(">>> st.secrets:", st.secrets)
import pandas as pd
import calendar
from datetime import datetime
from io import BytesIO

# manipulação de Word/Excel
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from openpyxl import load_workbook
from openpyxl.styles import PatternFill

# Airtable client
from pyairtable import Table



# --- Airtable setup ---

@st.cache_resource
def get_airtable_table(table_name: str):
    api_key = st.secrets["airtable_api_key"]
    base_id = st.secrets["airtable_base_id"]
    return Table(api_key, base_id, table_name)

tb_professores = get_airtable_table("professores")
tb_turmas      = get_airtable_table("turmas")
tb_horarios    = get_airtable_table("horarios")
tb_extras      = get_airtable_table("extras")

# --- Helpers de documento ---

map_hor = {
    "1ª": "7:00–7:50", "2ª": "7:50–8:40", "3ª": "8:40–9:30",
    "4ª": "9:50–10:40", "5ª": "10:40–11:30", "6ª": "12:20–13:10", "7ª": "13:10–14:00"
}
meses = ["Janeiro","Fevereiro","Março","Abril","Maio","Junho",
         "Julho","Agosto","Setembro","Outubro","Novembro","Dezembro"]
ano_planej = 2025

def extrai_serie(turma: str) -> str:
    return turma[:-1]

def set_border(par: Paragraph):
    p = par._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bd = OxmlElement('w:bottom')
    bd.set(qn('w:val'),'single'); bd.set(qn('w:sz'),'4')
    bd.set(qn('w:space'),'1'); bd.set(qn('w:color'),'auto')
    pBdr.append(bd); pPr.append(pBdr)

def insert_after(par: Paragraph, text='') -> Paragraph:
    new_p = OxmlElement('w:p'); par._p.addnext(new_p)
    para = Paragraph(new_p, par._parent)
    if text: para.add_run(text)
    return para

def gerar_agenda_template(entries, df_bank, professor, semana, bimestre, cores_turmas):
    wb = load_workbook("templates/agenda_modelo.xlsx")
    ws = wb.active
    ws["B1"] = professor
    ws["E1"] = semana
    row_map = {"1ª":4, "2ª":6, "3ª":8, "4ª":12, "5ª":14, "6ª":18, "7ª":20}
    col_map = {"Segunda":"C", "Terça":"D", "Quarta":"E", "Quinta":"F", "Sexta":"G"}
    for e in entries:
        col, row = col_map[e["dia"]], row_map[e["aula"]]
        ws[f"{col}{row}"] = f"{e['turma']} – {e['disciplina']}"
        fill = PatternFill(
            start_color=cores_turmas.get(e["turma"], "#FFFFFF").lstrip("#"),
            end_color=cores_turmas.get(e["turma"], "#FFFFFF").lstrip("#"),
            fill_type="solid"
        )
        ws[f"{col}{row}"].fill = fill
        ws[f"{col}{row+1}"] = (
            f"Aula {e['num']} – " +
            df_bank.loc[
                (df_bank["DISCIPLINA"]==e["disciplina"]) &
                (df_bank["ANO/SÉRIE"]==extrai_serie(e["turma"])) &
                (df_bank["BIMESTRE"]==bimestre) &
                (df_bank["Nº da aula"]==e["num"])
            ]["TÍTULO DA AULA"].iloc[0]
            if not df_bank.empty else ""
        )
        ws[f"{col}{row+1}"].fill = fill
    out = BytesIO(); wb.save(out); out.seek(0)
    return out

def gerar_plano_template(entries, df_bank, professor, semana, bimestre, turma,
                         metodologias, recursos, criterios, modelo="templates/modelo_plano.docx"):
    doc = Document(modelo)
    header_disciplinas = ", ".join(sorted({e['disciplina'] for e in entries}))
    total_aulas = str(len(entries))
    # Cabeçalho
    for p in doc.paragraphs:
        p.text = (p.text
                  .replace("ppp", professor)
                  .replace("ttt", turma)
                  .replace("sss", semana)
                  .replace("ddd", header_disciplinas)
                  .replace("nnn", total_aulas))
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.text = (p.text
                              .replace("ppp", professor)
                              .replace("ttt", turma)
                              .replace("sss", semana)
                              .replace("ddd", header_disciplinas)
                              .replace("nnn", total_aulas))
    # Blocos de aula + extras
    for p in doc.paragraphs:
        if p.text.strip() == "ccc":
            p.text = ""; last = p
            b0 = insert_after(last); set_border(b0); last = b0
            last = insert_after(last, "")
            for e in entries:
                sub = df_bank.loc[
                    (df_bank["DISCIPLINA"]==e["disciplina"]) &
                    (df_bank["ANO/SÉRIE"]==extrai_serie(turma)) &
                    (df_bank["BIMESTRE"]==bimestre) &
                    (df_bank["Nº da aula"]==e["num"])
                ]
                titulo = sub["TÍTULO DA AULA"].iloc[0] if not sub.empty else ""
                hab    = sub["HABILIDADE"].iloc[0]        if not sub.empty else ""
                cnt    = sub["CONTEÚDO"].iloc[0]         if not sub.empty else ""
                pa = insert_after(last, f"Aula {e['num']} – {titulo}"); pa.runs[0].bold=True; last=pa
                last = insert_after(last, "")
                ph = insert_after(last); rh=ph.add_run("Habilidade: "); rh.underline=True; ph.add_run(hab); last=ph
                last = insert_after(last, "")
                pc = insert_after(last); rc=pc.add_run("Conteúdo: "); rc.underline=True; pc.add_run(cnt); last=pc
                last = insert_after(last, "")
                b1 = insert_after(last); set_border(b1); last=b1
                last = insert_after(last, "")
            if metodologias:
                pm = insert_after(last); pm.add_run("Metodologia:").bold=True; last=pm
                for m in metodologias: last=insert_after(last, f"• {m}")
                last=insert_after(last, "")
            if recursos:
                pr = insert_after(last); pr.add_run("Recursos:").bold=True; last=pr
                for r in recursos: last=insert_after(last, f"• {r}")
                last=insert_after(last, "")
            if criterios:
                pc2 = insert_after(last); pc2.add_run("Critérios de Avaliação:").bold=True; last=pc2
                for c in criterios: last=insert_after(last, f"• {c}")
            break
    out = BytesIO(); doc.save(out); out.seek(0)
    return out

def gerar_guia_template(professor, turma, disciplina, bimestre, inicio, fim,
                        qtd_bimestre, qtd_semanal, metodologias, criterios,
                        df_bank, modelo="templates/modelo_guia.docx"):
    doc = Document(modelo)
    reps = {
        'ppp': professor,
        'ttt': turma,
        'bbb': bimestre,
        'iii': inicio.strftime('%d/%m/%Y'),
        'fff': fim.strftime('%d/%m/%Y'),
        'qqq': str(qtd_bimestre),
        'sss': str(qtd_semanal),
        'mmm': ", ".join(metodologias),
        'ccc': ", ".join(criterios),
        'ddd': disciplina
    }
    for p in doc.paragraphs:
        for k,v in reps.items():
            if k in p.text: p.text = p.text.replace(k,v)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for k,v in reps.items():
                        if k in p.text: p.text = p.text.replace(k,v)
    mask = (
        (df_bank["DISCIPLINA"]==disciplina)&
        (df_bank["ANO/SÉRIE"]==extrai_serie(turma))&
        (df_bank["BIMESTRE"]==bimestre)
    )
    habs = df_bank.loc[mask, "HABILIDADE"].dropna().astype(str).tolist()
    objs = df_bank.loc[mask, "OBJETO DE CONHECIMENTO"].dropna().astype(str).tolist()
    unique_habs = list(dict.fromkeys(habs))
    unique_objs = list(dict.fromkeys(objs))
    for p in doc.paragraphs:
        if 'hhh' in p.text: p.text = "\n".join(unique_habs)
        if 'ooo' in p.text: p.text = "\n".join(unique_objs)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if 'hhh' in p.text: p.text = "\n".join(unique_habs)
                    if 'ooo' in p.text: p.text = "\n".join(unique_objs)
    out = BytesIO(); doc.save(out); out.seek(0)
    return out

def gerar_planejamento_template(professor, disciplina, turma, bimestre,
                                grupos, df_bank, modelo="templates/modelo_planejamento.docx"):
    doc = Document(modelo)
    hdr = {'ppp':professor,'ddd':disciplina,'ttt':turma,'bbb':bimestre}
    for p in doc.paragraphs:
        for k,v in hdr.items():
            if k in p.text: p.text = p.text.replace(k,v)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for k,v in hdr.items():
                        if k in p.text: p.text = p.text.replace(k,v)
    for grp in grupos:
        p0 = doc.add_paragraph(); set_border(p0)
        doc.add_paragraph()
        doc.add_paragraph(f"Semana: {grp['semana']}")
        doc.add_paragraph(f"Aulas previstas: {grp['prev']}")
        doc.add_paragraph("Aulas dadas:")
        doc.add_paragraph("Aulas do material digital:")
        for n in grp['nums']:
            mask = (
                (df_bank["DISCIPLINA"]==disciplina)&
                (df_bank["ANO/SÉRIE"]==extrai_serie(turma))&
                (df_bank["BIMESTRE"]==bimestre)&
                (df_bank["Nº da aula"]==n)
            )
            titles = df_bank.loc[mask,"TÍTULO DA AULA"].dropna().astype(str).tolist()
            title = titles[0] if titles else ""
            doc.add_paragraph(f"Aula {n} – {title}")
        doc.add_paragraph("Objetivos:")
        objs = df_bank.loc[mask,"OBJETIVOS"].dropna().astype(str).tolist()
        for o in dict.fromkeys(objs): doc.add_paragraph(o)
        doc.add_paragraph("Habilidade:")
        habs = df_bank.loc[mask,"HABILIDADE"].dropna().astype(str).tolist()
        for h in dict.fromkeys(habs): doc.add_paragraph(h)
        doc.add_paragraph(f"Metodologia: {', '.join(grp['met'])}")
        doc.add_paragraph(f"Critérios de avaliação: {', '.join(grp['crit'])}")
        doc.add_paragraph()
    out = BytesIO(); doc.save(out); out.seek(0)
    return out

# --- Estado via Airtable ---

if "pagina" not in st.session_state:
    st.session_state.pagina = "Cadastro de Professor"

# carrega professores
if "professores" not in st.session_state:
    recs = tb_professores.all()
    st.session_state.professores = [
        {
            "id": r["id"],
            "email": r["fields"].get("email",""),
            "nome": r["fields"].get("nome",""),
            "disciplinas": r["fields"].get("disciplinas","").split(",")
        }
        for r in recs
    ]

# carrega turmas
if "turmas" not in st.session_state:
    recs = tb_turmas.all()
    # map email→list de turmas+cor
    turm_map = {}
    for r in recs:
        f = r["fields"]
        turm_map.setdefault(f["email"], []).append({
            "id": r["id"], "turma":f["turma"], "cor":f["cor"]
        })
    st.session_state.turmas = turm_map

# carrega horarios
if "horarios" not in st.session_state:
    recs = tb_horarios.all()
    st.session_state.horarios = [
        {
            **r["fields"],
            "id": r["id"]
        }
        for r in recs
    ]

# carrega extras
if "extras" not in st.session_state:
    recs = tb_extras.all()
    st.session_state.extras = {
        "metodologia": [r["fields"]["metodologia"] for r in recs if r["fields"].get("metodologia")],
        "recursos":     [r["fields"]["recurso"]     for r in recs if r["fields"].get("recurso")],
        "criterios":    [r["fields"]["criterio"]    for r in recs if r["fields"].get("criterio")]
    }

# --- Sidebar e navegação ---

pages = [
    "Cadastro de Professor","Cadastro de Turmas","Cadastro de Horário",
    "Gerar Agenda e Plano","Cadastro Extras","Gerar Guia",
    "Gerar Planejamento Bimestral"
]
for p in pages:
    if st.sidebar.button(p, use_container_width=True):
        st.session_state.pagina = p
    st.sidebar.markdown("\n")

# --- Páginas ---

# 1. Cadastro de Professor
if st.session_state.pagina == "Cadastro de Professor":
    st.header("Cadastro de Professor")
    nome = st.text_input("Nome")
    email = st.text_input("E-mail")
    disciplinas = st.multiselect(
        "Disciplina(s)",
        ["Arte","Ciências","Ed. Física","Ed. Financeira","Geografia","História",
         "Português","Inglês","Matemática","PV","Redação","Tecnologia","OE Port","OE Mat"]
    )
    if st.button("Salvar Professor"):
        # grava no Airtable
        rec = tb_professores.create({
            "email": email,
            "nome": nome,
            "disciplinas": ",".join(disciplinas),
            "senha_hash": ""  # ajustar login depois
        })
        st.experimental_rerun()
    # exibe
    for p in st.session_state.professores:
        st.write(f"{p['email']} — {p['nome']} — {', '.join(p['disciplinas'])}")

# 2. Cadastro de Turmas
elif st.session_state.pagina == "Cadastro de Turmas":
    st.header("Cadastro de Turmas")
    # lista de professores para escolher e-mail
    emails = [p["email"] for p in st.session_state.professores]
    prof_sel = st.selectbox("Professor (e-mail)", emails)
    turma = st.text_input("Turma (ex: 6ºA)")
    cor   = st.color_picker("Cor da Turma", "#FFFFFF")
    if st.button("Salvar Turma"):
        tb_turmas.create({
            "email": prof_sel,
            "turma": turma,
            "cor": cor.lstrip("#")
        })
        st.experimental_rerun()
    # exibe turmas do prof selecionado
    for t in st.session_state.turmas.get(prof_sel, []):
        st.write(f"{t['turma']} — #{t['cor']}")

# 3. Cadastro de Horário
elif st.session_state.pagina == "Cadastro de Horário":
    st.header("Cadastro de Horário")
    emails = [p["email"] for p in st.session_state.professores]
    prof_sel = st.selectbox("Professor (e-mail)", emails)
    turmas_prof = [t["turma"] for t in st.session_state.turmas.get(prof_sel,[])]
    turma_sel   = st.selectbox("Turma", turmas_prof)
    disciplina  = st.text_input("Disciplina")
    dia         = st.selectbox("Dia", ["Segunda","Terça","Quarta","Quinta","Sexta"])
    aula        = st.selectbox("Aula", list(map_hor.keys()))
    if st.button("Salvar Horário"):
        tb_horarios.create({
            "email": prof_sel,
            "turma": turma_sel,
            "disciplina": disciplina,
            "dia": dia,
            "aula": aula
        })
        st.experimental_rerun()
    # exibe tabela
    df_h = pd.DataFrame(st.session_state.horarios)
    st.dataframe(df_h)

# 4. Gerar Agenda e Plano
elif st.session_state.pagina == "Gerar Agenda e Plano":
    st.header("Gerar Agenda e Plano")
    if not st.session_state.horarios:
        st.warning("Cadastre horários primeiro.")
    else:
        df_bank = pd.read_excel("ES_banco.xlsx", header=0)
        prof_email = st.selectbox("Professor (e-mail)", [p["email"] for p in st.session_state.professores])
        prof_nome  = next(p["nome"] for p in st.session_state.professores if p["email"]==prof_email)
        bim = st.selectbox("Bimestre", ["1º","2º","3º","4º"])
        mes_nome  = st.selectbox("Mês", meses)
        semanas = [
            f"{w[0].strftime('%d/%m')} – {w[-1].strftime('%d/%m')}"
            for w in calendar.Calendar().monthdatescalendar(datetime.now().year, meses.index(mes_nome)+1)
            if w[0].month == meses.index(mes_nome)+1
        ]
        sem_sel = st.selectbox("Semana", semanas)
        cores_turmas = {t["turma"]:t["cor"] for t in st.session_state.turmas.get(prof_email,[])}
        entries = []
        for h in st.session_state.horarios:
            if h["email"]==prof_email:
                entries.append(h)
        if st.button("Gerar Agenda"):
            ag = gerar_agenda_template(entries, df_bank, prof_nome, sem_sel, bim, cores_turmas)
            st.download_button("Download Agenda", data=ag,
                               file_name="agenda_preenchida.xlsx",
                               mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        if st.button("Gerar Plano"):
            pl = gerar_plano_template(entries, df_bank, prof_nome, sem_sel, bim, prof_email,
                                      metodologias=st.session_state.extras["metodologia"],
                                      recursos=st.session_state.extras["recursos"],
                                      criterios=st.session_state.extras["criterios"])
            st.download_button("Download Plano", data=pl,
                               file_name=f"plano_{prof_email}.docx",
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# 5. Cadastro Extras
elif st.session_state.pagina == "Cadastro Extras":
    st.header("Cadastro Extras")
    sub = st.selectbox("Tipo", ["metodologia","recurso","criterio"])
    texto = st.text_input("Texto")
    if st.button("Inserir"):
        tb_extras.create({"email": "", sub: texto})
        st.experimental_rerun()
    st.write(st.session_state.extras)

# 6. Gerar Guia
elif st.session_state.pagina == "Gerar Guia":
    st.header("Gerar Guia")
    st.info("Em desenvolvimento…")

# 7. Gerar Planejamento Bimestral
elif st.session_state.pagina == "Gerar Planejamento Bimestral":
    st.header("Gerar Planejamento Bimestral")
    st.info("Em desenvolvimento…")
